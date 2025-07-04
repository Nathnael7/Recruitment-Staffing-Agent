import os
import fitz    # PyMuPDF
import docx
import textract
import spacy
from services.drive import download_file

import logging

# Setup logging
log_dir = ".logs"
os.makedirs(log_dir, exist_ok=True)
log_path = os.path.join(log_dir, "agents.log")
logging.basicConfig(
    filename=log_path,
    level=logging.INFO,
    format='%(asctime)s [%(levelname)s] %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)
nlp = spacy.load("en_core_web_sm")

def parse_resume(state):
    os.makedirs("tmp", exist_ok=True)
    parsed = []

    logging.info("Starting resume parsing for %d files.", len(state.get("resumes", [])))

    for file in state.get("resumes", []):
        fid, name = file.get("id"), file.get("name")
        path = f"tmp/{name}"
        logging.info("Downloading file: %s", name)
        try:
            download_file(fid, path)
        except Exception as e:
            logging.error("Failed to download file '%s': %s", name, e, exc_info=True)
            continue

        text = ""
        try:
            if name.endswith(".pdf"):
                logging.info("Parsing PDF: %s", name)
                with fitz.open(path) as d:
                    text = "\n".join(p.get_text() for p in d)

            elif name.endswith(".docx"):
                logging.info("Parsing DOCX: %s", name)
                d = docx.Document(path)
                text = "\n".join(p.text for p in d.paragraphs)
                for tbl in d.tables:
                    for row in tbl.rows:
                        for cell in row.cells:
                            text += "\n" + cell.text

            elif name.endswith(".doc"):
                logging.info("Parsing DOC: %s", name)
                text = textract.process(path).decode("utf-8")

            else:
                logging.warning("Unsupported file type: %s", name)
                continue

        except Exception as e:
            logging.error("Error parsing file '%s': %s", name, e, exc_info=True)
            continue

        try:
            doc_nlp = nlp(text)
            skills = [ent.text for ent in doc_nlp.ents if ent.label_ == "SKILL"]
        except Exception as e:
            logging.error("Error extracting skills from '%s': %s", name, e, exc_info=True)
            skills = []

        parsed.append({
            "name": name,
            "text": text,
            "skills": skills
        })
        logging.info("Finished processing file: %s", name)

    logging.info("Resume parsing completed. Parsed %d files.", len(parsed))
    return {"parsed_resumes": parsed}
